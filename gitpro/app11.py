from flask import Flask, request, jsonify, redirect, url_for, flash, session
import pandas as pd
import docx2txt
import PyPDF2
import pytesseract
import re
import nltk
from nltk.stem import WordNetLemmatizer
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from flask import render_template
from jinja2 import Template
import os
import docx
import shutil
from docx import Document
import transformers
from werkzeug.utils import secure_filename
from werkzeug.datastructures import FileStorage
import pathlib
import torch
from transformers import BertTokenizer, BertModel
from sklearn.metrics.pairwise import cosine_similarity
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import spacy
import numpy as np
from spacy.matcher import PhraseMatcher
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
from nltk.tokenize import sent_tokenize
from sentence_transformers import SentenceTransformer, util
from PIL import Image
import docx2txt
from PyPDF2 import PdfReader
import moviepy
from moviepy import VideoFileClip
from moviepy import AudioFileClip
import speech_recognition as sr
import librosa
from deepface import DeepFace
from transformers import pipeline
from transformers import Wav2Vec2ForCTC, Wav2Vec2Processor
from textblob import TextBlob
import cv2
import openai
nltk.download('wordnet')
nltk.download('punkt')
nltk.download('stopwords')
nlp=spacy.load('en_core_web_sm')
app = Flask(__name__)

# Define the upload folders
RESUME_FOLDER = os.path.join(os.getcwd(), 'uploads', 'resumes')
JOB_DESCRIPTION_FOLDER = os.path.join(os.getcwd(), 'uploads', 'job_descriptions')

# Ensure directories exist
os.makedirs(RESUME_FOLDER, exist_ok=True)
os.makedirs(JOB_DESCRIPTION_FOLDER, exist_ok=True)

#folder for video resume
UPLOAD_FOLDER1 = 'uploaded_videos'
ANALYSIS_FOLDER1= 'analysis_results'
os.makedirs(UPLOAD_FOLDER1, exist_ok=True)
os.makedirs(ANALYSIS_FOLDER1, exist_ok=True)
app.config['UPLOAD_FOLDER1'] = UPLOAD_FOLDER1

# Allowed extensions for video files
ALLOWED_EXTENSIONS1 = {'mp4', 'avi', 'mov', 'mkv', 'flv', 'wmv'}

app.secret_key = 'finalyearproject'
tokenizer1 = AutoTokenizer.from_pretrained('valhalla/t5-base-qg-hl')
model1 = AutoModelForSeq2SeqLM.from_pretrained('valhalla/t5-base-qg-hl')
# Allowed file extensions
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'txt', 'csv','png','jpg','jpeg'}
EMAIL_ADDRESS = "ramiavv39@gmail.com"
EMAIL_PASSWORD = "ramya@555"
tokenizer = BertTokenizer.from_pretrained('bert-base-uncased')
model = BertModel.from_pretrained('bert-base-uncased')
model.eval()
bert_model = SentenceTransformer('paraphrase-MiniLM-L6-v2')
processor = Wav2Vec2Processor.from_pretrained("facebook/wav2vec2-base-960h")
model1 = Wav2Vec2ForCTC.from_pretrained("facebook/wav2vec2-base-960h")
technical_skills= [
    # Programming Languages
    "Python", "Java", "C++", "C", "C#", "JavaScript", "HTML", "CSS", "TypeScript", 
    "Swift", "Kotlin", "Ruby", "Rust", "Go (Golang)", "PHP", "R", "Scala", "Perl", "MATLAB",

    # Frameworks & Libraries
    "Django", "Flask", "Spring", "React", "Angular", "Vue.js", "Node.js", "Express.js", 
    "Ruby on Rails", "ASP.NET", "JUnit", "TensorFlow", "PyTorch", "Keras", "Scikit-learn", 
    "Pandas", "NumPy", "OpenCV",

    # Cloud Computing & DevOps
    "Amazon Web Services (AWS)", "Microsoft Azure", "Google Cloud Platform (GCP)", 
    "Kubernetes", "Docker", "Jenkins", "Terraform", "Ansible", "OpenShift", 
    "Vagrant", "Puppet", "Chef", "CloudFormation", "Serverless architecture", 
    "CI/CD pipelines",

    # Databases & Data Management
    "SQL", "NoSQL", "MySQL", "PostgreSQL", "MongoDB", "Cassandra", "Couchbase", 
    "Redis", "DynamoDB", "Elasticsearch", "Hadoop", "Apache Spark", "Apache Kafka", 
    "HBase", "Hive", "Presto", "Snowflake", "Oracle", "MariaDB", "Firebase", 
    "BigQuery",

    # Data Science & Machine Learning
    "Machine Learning", "Deep Learning", "Natural Language Processing (NLP)", 
    "Computer Vision", "Data Analysis", "Data Mining", "Predictive Modeling", 
    "Reinforcement Learning", "Feature Engineering", "Model Optimization", 
    "Bayesian Networks", "Time Series Analysis", "AutoML", "Clustering (K-Means, DBSCAN)", 
    "Dimensionality Reduction (PCA, LDA)", "XGBoost", "LightGBM",

    # Embedded Systems & IoT
    "Embedded C", "ARM Cortex", "RTOS", "FPGA programming", "Microcontroller programming", 
    "IoT (Internet of Things)", "Arduino", "Raspberry Pi", "Embedded Linux", 
    "Real-Time Systems", "UART", "SPI", "I2C protocols", "CAN protocol", "VHDL", 
    "Verilog", "Bluetooth Low Energy (BLE)", "ZigBee", "LoRaWAN", "MQTT",

    # Networking & Security
    "TCP/IP", "DNS", "DHCP", "IPv6", "SDN", "Network Virtualization", "VPN", "VLAN", 
    "Firewall configuration", "Penetration Testing", "Ethical Hacking", "Wireshark", 
    "Network Forensics", "IDS/IPS", "SIEM", "SSL/TLS", "Zero Trust Security", 
    "OAuth", "SSH",

    # Operating Systems
    "Linux (Ubuntu, CentOS, RHEL)", "Windows", "macOS", "Unix", "FreeRTOS", "QNX", 
    "Android", "iOS", "Shell Scripting (Bash, PowerShell)", "Kernel Programming", 
    "Virtualization (VMware, Hyper-V, KVM)",

    # Version Control & Collaboration
    "Git", "GitHub", "GitLab", "Bitbucket", "Subversion (SVN)", "Mercurial", 
    "Jira", "Confluence", "Slack", "Trello",

    # Software Development & Design
    "Object-Oriented Programming (OOP)", "Functional Programming", "Test-Driven Development (TDD)", 
    "Behavior-Driven Development (BDD)", "Agile methodologies", "SCRUM", "Kanban", 
    "Design Patterns", "SOLID Principles", "UML Diagrams", "API Development (REST, GraphQL)", 
    "Microservices Architecture", "SOA", "Web Services (SOAP, REST)", "Message Queuing", 
    "Distributed Systems", "Event-Driven Architecture", "Software Testing (Unit, Integration, E2E)",

    # Engineering Tools & Simulation Software
    "AutoCAD", "MATLAB/Simulink", "CATIA", "SolidWorks", "ANSYS", "PSpice", 
    "LabVIEW", "COMSOL Multiphysics", "NI Multisim", "Fusion 360",

    # Robotics & Automation
    "ROS", "PLC Programming", "SCADA systems", "Drones and UAV development", 
    "Robot kinematics", "Control systems", "Industrial automation", 
    "Sensors and Actuators", "Simulink", "Path planning algorithms"
]

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Route to serve the HTML page
@app.route('/')
def home():
    return render_template('index5.html')

@app.route('/upload_combined', methods=['POST'])
def upload():
    # Check if files are part of the request
    print(request.files)
    print(request.content_type)
    if 'resumes' not in request.files or 'job_description' not in request.files:
        return 'No file part'

    resume_files = request.files.getlist('resumes')
    job_file = request.files['job_description']
    
    uploaded_filenames = []
    for file in resume_files:
        if file.filename == '':
            return 'No selected file for resumes'
        if file and allowed_file(file.filename):
            save_path = os.path.join(RESUME_FOLDER, file.filename)
            file.save(save_path)
            uploaded_filenames.append(save_path)

    session['uploaded_resumes'] = uploaded_filenames

    if job_file and allowed_file(job_file.filename):
        job_file_path = os.path.join(JOB_DESCRIPTION_FOLDER, job_file.filename)
        job_file.save(job_file_path)
        session['uploaded_job'] = job_file_path

    return process(resume_files, job_file)

def process(resume_files, job_file):
    if resume_files:
        uploaded_resumes_df = process_uploaded_files(resume_files)

    if job_file:
        processed_job_desc_df = process_uploaded_job_descriptions(job_file)
    filepath1=os.path.join('cleaned','uploaded_resumes_df.csv')
    uploaded_resumes_df.to_csv(filepath1,index=False)
    filepath2=os.path.join('cleaned','processed_job_desc_df.csv')
    processed_job_desc_df.to_csv(filepath2,index=False)
        # Process top-ranked resumes
        #top_n_resumes = int(request.form.get('top_n_resumes', 1))  # Default to 1 if not provided
        #ranked_resumes, similarity_scores = rank_resumes(uploaded_resumes_df,processed_job_desc_df, tokenizer, model)
        # Get the number of resumes to save from the HR input
        #num_resumes = int(request.form.get('numTopResumes'))
        # Save the top-ranked resumes
        #save_top_resumes(ranked_resumes, num_resumes, "top_resumes")
        #ranked_resumes = rank_resumes(uploaded_resumes_df, processed_job_desc_df,tokenizer,model)
        #save_interview_questions(ranked_resumes, num_resumes, "interview_questions")
        #interview_questions = generate_interview_questions(ranked_resumes)
        #send_feedback_to_candidates(ranked_resumes)
        #send_feedback(resume_df, job_desc_df, ranked_resumes, similarity_scores)   
    return "Submitted Successfully"

@app.route('/top-ranked-resumes',methods=['Post'])
def topu():
    data=request.get_json()
    num_resumes = data.get('numTopResumes')
    filepath1=os.path.join('cleaned','uploaded_resumes_df.csv')
    filepath2=os.path.join('cleaned','processed_job_desc_df.csv')
    uploaded_resumes_df=pd.read_csv(filepath1)
    processed_job_desc_df=pd.read_csv(filepath2)
    print(uploaded_resumes_df)
    print(processed_job_desc_df)
    from docx import Document
    doc = Document()
    doc.add_paragraph(num_resumes)
    doc.save('num_resume.docx')
    ranked_resumes, similarity_scores = rank_resumes(uploaded_resumes_df,processed_job_desc_df, tokenizer, model)
        # Get the number of resumes to save from the HR input
    
        # Save the top-ranked resumes
    filepath3=os.path.join('cleaned','ranked_resumes_df.csv')
    filepath4=os.path.join('cleaned','similarity_score.csv')
    ranked_resumes.to_csv(filepath3,index=False)
    similarity_scores_df = pd.DataFrame(similarity_scores)  # Assuming similarity_scores is a NumPy array
    similarity_scores_df.to_csv(filepath4, index=False)
    print("THE RANKED RESUMES ARE")
    print(ranked_resumes)
    save_top_resumes(ranked_resumes, num_resumes, "top_resumes")
    print("Reached return statement")
    return "Top ranked resumes are stored successfully"
@app.route('/generate-interview-questions',methods=['POST'])
def interv():
    filepath3=os.path.join('cleaned','ranked_resumes_df.csv')
    ranked_resumes=pd.read_csv(filepath3)
    save_interview_questions(ranked_resumes,"interview_questions")
    return "Inteview questions stored succesfully"
@app.route('/send-feedback',methods=['Post'])
def feedb():
    filepath1=os.path.join('cleaned','uploaded_resumes_df.csv')
    filepath2=os.path.join('cleaned','processed_job_desc_df.csv')
    resumes_df=pd.read_csv(filepath1)
    job_desc_df=pd.read_csv(filepath2)
    filepath3=os.path.join('cleaned','ranked_resumes_df.csv')
    filepath4=os.path.join('cleaned','similarity_score.csv')
    ranked_resumes=pd.read_csv(filepath3)
    similarity_scores=pd.read_csv(filepath4)
    send_feedback(resumes_df, job_desc_df, ranked_resumes, similarity_scores)
    return "Feedback Sent Successfully"
    
# Function to extract text from different file types
def extract_text_from_file(file_path):
  filename = os.path.basename(file_path)
  file_extension = os.path.splitext(filename)[1].lower()

  if file_extension == '.txt':
    with open(file_path, 'rb') as file:
      text = file.read().decode('utf-8')
  elif file_extension == '.docx':
    with open(file_path, 'rb') as file:
      text = docx2txt.process(file)
  elif file_extension == '.pdf':
    with open(file_path, 'rb') as file:
      reader = PdfReader(file)
      text = ""
      for page in reader.pages:
        text += page.extract_text()
  elif file_extension in ('.jpg', '.jpeg', '.png'):
    # Text extraction for image files
    img = Image.open(file_path)
    text = pytesseract.image_to_string(img)
  else:
    raise ValueError("Unsupported file format. Please upload a .txt, .docx, .pdf, .jpg, .jpeg, or .png file.")
  return text
def get_bert_embeddings(texts, tokenizer, model):
    embeddings = []
    for text in texts:
        if not isinstance(text, str):
            raise ValueError("Each input text must be a string.")
        inputs = tokenizer(text, return_tensors='pt', max_length=512, truncation=True, padding=True)
        with torch.no_grad():
            outputs = model(**inputs)
        last_hidden_states = outputs.last_hidden_state
        cls_embedding = last_hidden_states.mean(dim=1).squeeze().numpy()
        embeddings.append(cls_embedding)
    return embeddings
def rank_resumes(resume_df, job_desc_df, tokenizer, model):
    resume_texts = resume_df['Cleaned_Text'].astype(str).tolist()
    job_description_text = job_desc_df['Cleaned_Description'].astype(str).iloc[0]
    resume_embeddings = get_bert_embeddings(resume_texts, tokenizer, model)
    job_desc_embedding = get_bert_embeddings([job_description_text], tokenizer, model)[0]
    print("the embeddings are")
    print(resume_embeddings)
    print(job_desc_embedding)
    similarities = cosine_similarity([job_desc_embedding], resume_embeddings).flatten()
    print("the similarities are")
    print(similarities)
    ranked_indices = similarities.argsort()[::-1]
    ranked_resumes = resume_df.iloc[ranked_indices]
    #ranked_resumes=resume_df.iloc['skills']
    #ranked_resumes=resume_df.iloc['qualifications']
    # Add a new column "RESUME_ID" with the ranked indices
    ranked_resumes['Resume_ID'] = ranked_resumes.index.values[ranked_indices]
    ranked_resumes = ranked_resumes[['Resume_ID', 'Cleaned_Text', 'skills', 'qualifications','Name','Email','filename']]
    return ranked_resumes, similarities[ranked_indices]

# Rank resumes
#ranked_resumes, similarity_scores = rank_resumes(resume_df, job_desc_df, tokenizer, model)

# Output ranked resumes
#for idx, (index, row) in enumerate(ranked_resumes.iterrows()):
    #print(f"Rank {idx + 1}: Resume ID {row.get('Resume_ID', 'N/A')} (Similarity Score: {similarity_scores[idx]:.4f})")
    #print(f"Resume Text: {row['Cleaned_Text']}\n")

# Save top-ranked resumes to a local folder
def save_top_resumes(ranked_resumes, num_resumes, output_folder):
    print("the ranked resumes columns are: ")
    print(ranked_resumes.columns)
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    for idx, (index, row) in enumerate(ranked_resumes.iterrows()):
        if idx < int(num_resumes):
            # Construct the source file path from 'upload/resume' folder
            source_file_path = os.path.join("uploads/resumes", row['filename'])
            destination_file_path = os.path.join(output_folder, row['filename'])    
            shutil.copy(source_file_path, destination_file_path)
            print(f" {row['filename']} to {output_folder}")

    
    

# Get the number of resumes to save from the HR input
#num_resumes = int(request.form.get('numTopResumes'))
# Save the top-ranked resumes
#save_top_resumes(ranked_resumes, num_resumes, "top_resumes")
# Define a function to generate interview questions for a given resume
def generate_interview_questions(resume_text):
    # This is a simple example of generating interview questions based on the resume text
    # You can replace this with your own logic to generate interview questions
    print(resume_text)
    questions = []
    for sentence in resume_text.split('.'):
        if 'experience' in sentence.lower():
            questions.append('Can you tell me more about your experience in ' + sentence.split('experience')[1].strip() + '?')
        elif 'skills' in sentence.lower():
            questions.append('How do you think your skills in ' + sentence.split('skills')[1].strip() + ' will help you in this role?')
    return questions



# Function to generate questions from a sentence
def generate_question(sentence):
    # Add highlight tokens <hl> around the sentence
    input_text = f"generate question: <hl> {sentence} <hl>"    
    # Tokenize input text
    inputs = tokenizer1.encode(input_text, return_tensors='pt', max_length=512, truncation=True)
    # Generate question
    outputs = model1.generate(inputs, max_length=100, num_beams=5, early_stopping=True) 
    # Decode and return the generated question
    return tokenizer1.decode(outputs[0], skip_special_tokens=True)
# Function to generate and return questions from the paragraph
def get_generated_questions(paragraph):
    # Tokenize the paragraph into sentences
    sentences = sent_tokenize(paragraph)
    # Generate questions from each sentence
    return [generate_question(sentence) for sentence in sentences]
# Get the generated questions

# Return the list of generated questions

# Define a function to save interview questions to a local folder
def save_interview_questions(ranked_resumes,output_folder):
    print("THE COLUMNS ARE:")
    print(ranked_resumes.columns)
    doc = Document('num_resume.docx')
    for paragraph in doc.paragraphs:
        num_resumes=paragraph.text
        
    print(num_resumes)
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    for idx, (index, row) in enumerate(ranked_resumes.iterrows()):
        if idx < int(num_resumes):
            #sentences = sent_tokenize(paragraph)
            #questions = generate_interview_questions(row['Cleaned_Text'])
            questions = get_generated_questions(row['Cleaned_Text'])
            document = Document()
            document.add_heading('Interview Questions for ' + str(row['Resume_ID']), 0)
            for question in questions:
                document.add_paragraph(question)
            document.save(os.path.join(output_folder, str(row['Resume_ID']) + '.docx'))

# Get the number of resumes to save from the HR input


# Save the interview questions
#save_interview_questions(ranked_resumes, num_resumes, "interview_questions")
def generate_feedback(resume_df, job_description_df):
    #print("the resume is")
    #print(resume_df)
    #print(type(resume_df))
    import ast
    parsed_job_skills = ast.literal_eval(job_description_df)# Extract and parse the job description set
    parsed_resume_skills = ast.literal_eval(resume_df)  # Extract and parse a single resume's skill set

    print(parsed_resume_skills)
    missing_skills = parsed_job_skills - parsed_resume_skills
  
    if missing_skills:
        feedback = f"{', '.join(missing_skills)}."
    else:
        feedback = "Your technical skills match the job description well."  
    return feedback
# Define a function to send feedback via email
def send_feedback(resume_df, job_desc_df, ranked_resumes, similarity_scores):
    doc = Document('num_resume.docx')
    for paragraph in doc.paragraphs:
        num_resumes=paragraph.text
    print(num_resumes)
    print(ranked_resumes.columns)
    for idx, (index, row) in enumerate(ranked_resumes.iterrows()):
        print(idx)
        if idx+1 > int(num_resumes):  # Send feedback to non-top 10 resumes
            print("nice")
            #resume_skills = row['skills']
            #resume_qualifications = row['qualifications']
            #print(resume_skills)
            #print(resume_qualifications)
            #resume_combined_text = resume_skills + " " + resume_qualifications
            
            # Extract skills and qualifications from the job description dataframe
            #job_desc_skills = job_desc_df['skills'].iloc[0]
            #job_desc_qualifications = job_desc_df['qualifications'].iloc[0]
            #job_desc_combined_text = job_desc_skills + " " + job_desc_qualifications
            
            # Generate feedback
            #resume_skills=extract_technical_skills(resume_df['Cleaned_Text'].to_string(),technical_skills)
            #job_description_skills=extract_technical_skills(job_desc_df['Cleaned_Description'].to_string(),technical_skills)
            resume_skills=row['skills']
            job_description_skills=job_desc_df['skills'].iloc[0]
            #print(type(resume_df['Cleaned_Text'].to_string()))
            #print(type(job_desc_df['Cleaned_Description'].to_string()))
            print("The resume skills are: ")
            print(resume_skills)
            print("the job description skills are")
            print(job_description_skills)
            feedback = generate_feedback(resume_skills,job_description_skills)
            #print(feedback)
            msg = MIMEMultipart()
            msg['From'] = 'ramiavv39@gmail.com'
            msg['To'] = row['Email']
            msg['Subject'] = 'Feedback on your resume'
            body = 'Dear ' + 'candidate'+ ',\n\nWe appreciate your interest in the position. However, after reviewing your resume, we noticed that it lacks skills like '
            for item in feedback:
                body += item + ""
            body += '  As these skills are required for the role which lacks in your resume we are not able to proceed further with your candidature. We hope this feedback is helpful in your job search. Thank you for considering our company.\n\nBest regards,\nAli'
            msg.attach(MIMEText(body, 'plain'))
            server = smtplib.SMTP('smtp.gmail.com', 587)
            server.starttls()
            server.login(msg['From'], 'umyi lvse tcxp sowv')
            text = msg.as_string()
            print(text)
            print(msg['To'])
            server.sendmail(msg['From'], msg['To'], text)
            server.quit()

# Send feedback to non-top resumes
#send_feedback(resume_df, job_desc_df, ranked_resumes, similarity_scores)
# Function to clean text
def clean_text(text):
    stop_words = set(stopwords.words('english'))
    lemmatizer = nltk.WordNetLemmatizer()
    
    # 1. Remove email addresses, URLs, and phone numbers
    #text = re.sub(r'\S+@\S+', '', text)  # Remove emails
    text = re.sub(r'http\S+', '', text)  # Remove URLs
    text = re.sub(r'\b\d{10,}\b', '', text)  # Remove phone numbers
    # 2. Remove special characters, numbers, punctuation, and HTML tags
    text = re.sub(r'<.*?>', '', text)  # Remove HTML tags
    text = re.sub(r'[^A-Za-z\s]', '', text)  # Remove non-alphabetical characters
    #text = re.sub(r'\d+', '', text)  # Remove digits/numbers
    # 3. Convert to lowercase
    text = text.lower()
    # 4. Tokenize (split text into words)
    words = text.split()
    # 5. Remove stopwords
    words = [word for word in words if word not in stop_words]
    # 6. Lemmatization (reduce words to their root form)
    words = [lemmatizer.lemmatize(word) for word in words]
    # 7. Remove repeated words (optional)
    words = list(dict.fromkeys(words))
    # 8. Rejoin the words into a single string
    cleaned_text = ' '.join(words)
    # 9. Remove extra spaces
    cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()
    return cleaned_text


def extract_name_and_email(resume_text):
    #nlp_best=spacy.load('model-best')
    dok=nlp(resume_text)
    #name_pattern = r"([A-Z][a-z]+) ([A-Z][a-z]+)"
    email_pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b"
    names = [ent.text for ent in dok.ents if ent.label_ == 'PERSON']
    #name_match = re.search(name_pattern, resume_text)
    email_match = re.search(email_pattern, resume_text)

    if email_match:
        #name = name_match.group()
        email = email_match.group()
        print("the name is")
        print(names)
        print("the email is ")
        print(email)
        return names, email
    else:
        return None, None
def extract_skill_quali(resume_text):
    #nlp_best=spacy.load('model-best')
    dok=nlp(resume_text)
    skills=[]
    qualifications=[]
    matcher = PhraseMatcher(nlp.vocab)
    skills_list = [
    # Software Engineering Skills
    "Python", "Java", "C++", "C#", "JavaScript", "HTML/CSS", "SQL", "R", "MATLAB", "Ruby", "Swift", "PHP",
    "Go", "TypeScript", "Node.js", "Angular", "React", "Vue.js", "Django", "Flask", "Spring Boot", "Git",
    "Docker", "Kubernetes", "REST APIs", "GraphQL", "Machine Learning", "Deep Learning", "Artificial Intelligence",
    "Data Science", "Data Analysis", "TensorFlow", "PyTorch", "Computer Vision", "Natural Language Processing (NLP)",
    "Agile Methodology", "Scrum", "DevOps", "CI/CD", "Cloud Computing (AWS, Azure, GCP)", "Microservices",
    "Unit Testing", "TDD (Test-Driven Development)", "BDD (Behavior-Driven Development)",

    # Mechanical Engineering Skills
    "CAD (Computer-Aided Design)", "SolidWorks", "AutoCAD", "CATIA", "ANSYS", "FEA (Finite Element Analysis)",
    "CFD (Computational Fluid Dynamics)", "Thermodynamics", "Fluid Mechanics", "Heat Transfer", 
    "HVAC (Heating, Ventilation, and Air Conditioning)", "Mechanical Design", "Product Development", "Mechatronics",
    "Robotics", "PLC (Programmable Logic Controller)", "CNC Machining", "Manufacturing Processes", "Material Science",
    "Quality Control", "Lean Manufacturing", "Six Sigma", "Failure Mode and Effects Analysis (FMEA)", 
    "Project Management",

    # Electrical Engineering Skills
    "Circuit Design", "PCB Design", "VHDL/Verilog", "Embedded Systems", "Microcontrollers", 
    "FPGA (Field-Programmable Gate Array)", "Digital Signal Processing (DSP)", "Power Electronics",
    "Analog Circuit Design", "Control Systems", "MATLAB/Simulink", "LabVIEW", "SCADA (Supervisory Control and Data Acquisition)",
    "PLC Programming", "Electrical Wiring", "Electrical Safety", "Power Distribution", "Renewable Energy Systems",
    "Instrumentation", "Troubleshooting", "Wireless Communication", "RF Engineering", "Signal Processing",

    # Civil Engineering Skills
    "Structural Analysis", "AutoCAD Civil 3D", "Revit", "STAAD.Pro", "SAP2000", "ETABS", "Surveying", 
    "Geotechnical Engineering", "Construction Management", "Building Information Modeling (BIM)", "Quantity Surveying",
    "Project Scheduling (Primavera P6, MS Project)", "Environmental Impact Assessment", "Transportation Engineering",
    "Hydraulic Engineering", "Concrete Design", "Steel Design", "Soil Mechanics", "Urban Planning", 
    "Site Inspection", "Water Resource Management", "Drainage Design", "GIS (Geographic Information Systems)",
    "Estimation and Costing",

    # Chemical Engineering Skills
    "Process Simulation (Aspen HYSYS, Aspen Plus)", "Process Design", "Chemical Process Control", "Heat Exchangers",
    "Distillation", "Mass Transfer", "Fluid Dynamics", "Thermodynamics", "Chemical Reaction Engineering",
    "Polymer Engineering", "Catalysis", "Environmental Engineering", "Wastewater Treatment", "Safety Management",
    "Piping and Instrumentation Diagrams (P&ID)", "Plant Design", "Material Balances", "Process Optimization",
    "Quality Assurance", "Hazard Analysis (HAZOP)", "Laboratory Analysis", "Process Scale-up", 
    "Petrochemical Engineering",

    # General Engineering Skills
    "Project Management", "Technical Writing", "Data Analysis", "Root Cause Analysis", "Research & Development (R&D)",
    "Problem-Solving", "Critical Thinking", "Troubleshooting", "Lean Engineering", "Quality Assurance/Control",
    "Risk Assessment", "Communication Skills", "Team Collaboration", "Presentation Skills", "Time Management",
    "Leadership", "AutoCAD", "3D Printing", "Prototyping", "Simulation Modeling",

    # Emerging Technologies Skills
    "Internet of Things (IoT)", "Blockchain", "5G Technology", "Cybersecurity", "Big Data", "Cloud Computing",
    "Virtual Reality (VR) / Augmented Reality (AR)", "Edge Computing", "Robotics Process Automation (RPA)",
    "Quantum Computing", "Autonomous Systems", "Smart Grids", "Renewable Energy Systems"
]

    skill_patterns = [nlp(skill) for skill in skills_list]
    matcher.add("SKILLS", skill_patterns)
    skills = set([dok[start:end].text for match_id, start, end in matcher(dok)])
    # Extract qualifications
    qualification_keywords = ["Bachelor", "Master", "Ph.D.", "Diploma", "Certification"]
    qualifications = [ent.text for ent in dok.ents if any(keyword.lower() in ent.text.lower() for keyword in qualification_keywords)]
    print("the skills are")
    print(skills)
    print("the qualifications are")
    print(qualifications)
    return  skills,qualifications
    


def process_uploaded_files(files):
    resumes = []  
    for file in files:
        if not isinstance(file, FileStorage):
            # Create a FileStorage object from the file
            file = FileStorage(file, filename=os.path.basename(file))
        
        filename = secure_filename(file.filename)
        file_path=os.path.join(RESUME_FOLDER,filename)
        try:
            text = extract_text_from_file(file_path)
            Name,Email=extract_name_and_email(text)
            skills,qualifications=extract_skill_quali(text)
            cleaned_text = clean_text(text)
            resumes.append({'filename': filename, 'Cleaned_Text': text, 'Name':Name,'Email':Email,'skills': skills,'qualifications':qualifications})
        except Exception as e:
            print(f"Error processing file {filename}: {str(e)}")
    
    # Convert to DataFrame for further processing
    resumes_df = pd.DataFrame(resumes)
    print("the resume wth name and email are:")
    print(resumes_df.columns)
    print(resumes_df)
    return resumes_df

# Function to process uploaded files for job descriptions
def process_uploaded_job_descriptions(file):
    job_descriptions = []
    if not isinstance(file, FileStorage):
            file = FileStorage(file, filename=os.path.basename(file))
    filename = secure_filename(file.filename)
    file_path = pathlib.Path(JOB_DESCRIPTION_FOLDER)/filename
    print("File Path:", file_path)
    if os.path.exists(file_path):
        try:
            with open(file_path, 'rb') as file:
                print("nice")
        except Exception as e:
            print("Error opening file:", str(e))
    else:
        print("File does not exist:", file_path)
        print(os.listdir(JOB_DESCRIPTION_FOLDER))
        print("File name:", filename)
        print("Is file hidden?", os.path.isfile(file_path))
    try:
        text = extract_text_from_file(file_path)
        cleaned_text = clean_text(text)
        skills,qualifications=extract_skill_quali(text)
        job_descriptions.append({'filename': filename, 'Cleaned_Description': text,'skills':skills,'qualifications':qualifications})
    except Exception as e:
        print(f"Error processing file {filename}: {str(e)}")
    
    # Convert to DataFrame for further processing
    job_descriptions_df = pd.DataFrame(job_descriptions)
    print(job_descriptions_df)
    return job_descriptions_df
def gienerate_feedback(resume_skills, job_description_skills):
    print(resume_skills)
    print(job_description_skills)
    missing_skills = job_description_skills - resume_skills
    
    if missing_skills:
        feedback = f"Your resume is missing the following key technical skills: {', '.join(missing_skills)}."
    else:
        feedback = "Your technical skills match the job description well."
    
    return feedback

def extract_technical_skills(text, technical_skills):
    doc = nlp(text.lower())
    found_skills = set()
    normalized_skills = {skill.lower() for skill in technical_skills}
    print("the normalized skills are")
    print(normalized_skills)
    print("the doc is")
    print(doc)
    # Loop through each token in the text
    for token in doc:
        if token.text in normalized_skills:
            print(token.text)
            found_skills.add(token.text)
    
    return found_skills


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS1
# Step 1: Extract Audio and Frames for a Single Video
def extract_audio(video_path):
        audio_path = video_path.rsplit('.', 1)[0] + '.wav'
        video_clip=VideoFileClip(video_path)
        audio_clip=video_clip.audio
        print("the audio path is")
        print(audio_path)
        audio_clip.write_audiofile(audio_path)
        audio_clip.close()
        video_clip.close()
        print("Audio extraction successful")
        return audio_path
def split_audio(audio_path, segment_length=30):
    """
    Splits an audio file into smaller chunks of specified segment length (in seconds).
    """
    from pydub import AudioSegment
    from pydub.utils import make_chunks
    audio = AudioSegment.from_wav(audio_path)
    chunks = make_chunks(audio, segment_length * 1000)  # pydub works in milliseconds
    chunk_paths = []
    for i, chunk in enumerate(chunks):
        chunk_name = f"chunk_{i}.wav"
        chunk.export(chunk_name, format="wav")
        chunk_paths.append(chunk_name)
    return chunk_paths

def transcribe_audio_chunks(chunk_paths):
    
    """
    Transcribes a list of audio chunks and combines the results.
    """
    recognizer = sr.Recognizer()
    transcript = ""
    for chunk_path in chunk_paths:
        with sr.AudioFile(chunk_path) as source:
            audio_data = recognizer.record(source)
            try:
                text = recognizer.recognize_google(audio_data)
                transcript += text + " "
            except sr.UnknownValueError:
                transcript += "[Unrecognizable Speech] "
            except sr.RequestError as e:
                print(f"Error with Google Speech Recognition API: {e}")
        os.remove(chunk_path)  # Clean up the chunk file
    return transcript

def extract_resume_info(paragraph):
    
    doc = nlp(paragraph)
    resume_info = {
        "name": "",
        "education": [],
        "skills": [],
        "experience": []
        }
    name = 'ram'
    education_keywords = ["B.Tech", "M.Sc.", "Bachelor's", "Master's", "PhD", "University", "College", "Degree"]
    for sent in doc.sents:
        if any(keyword in sent.text for keyword in education_keywords):
            # Extract degree, major, and university (requires more sophisticated parsing)
            resume_info["education"].append({"degree": "", "major": "", "university": ""}) 
    skill_keywords = ["skills", "proficient", "expertise", "knowledge"]
    for token in doc:
        if token.pos_ == "NOUN" or token.pos_ == "ADJ":
            if any(keyword in token.text.lower() for keyword in skill_keywords):
                resume_info["skills"].append(token.text)
    # Extract Experience (Basic - needs refinement)
    experience_keywords = ["experience", "worked", "role", "responsibilities", "achievements"]
    for sent in doc.sents:
        if any(keyword in sent.text for keyword in experience_keywords):
            resume_info["experience"].append(sent.text)
    # Save the resume to a file
    filename = f"{name.replace(' ', '_')}_resume.txt"
    file_path = os.path.join('generated_resumes', filename)
    #with open(file_path, "w") as f:
        #f.write(resume_content)


# Example Usage


def process_text_files(folder_path):
    text_contents = []
    for filename in os.listdir(folder_path):
        if filename.endswith(".txt"):
            file_path = os.path.join(folder_path, filename)
            try:
                with open(file_path, 'r') as file:
                    text = file.read()
                    return text
            except FileNotFoundError:
                print(f"Error: File not found at {file_path}")
    return None

def process_video(video_path, output_text_path):
    """
    Processes a single video resume for speech-to-text, tone, fluency, and behavior analysis.

    Args:
        video_path (str): Path to the video file.
        output_text_path (str): Path to save the analysis result as a text file.
    """
    

    audio_path = extract_audio(video_path)
    print("the aduio is shown below")
    print(audio_path)
    chunk_paths = split_audio(audio_path, segment_length=30)  
    transcription = transcribe_audio_chunks(chunk_paths)
    print("transcription done successfully")
    
    # Analyze tone (sentiment)
    blob = TextBlob(transcription)
    sentiment = blob.sentiment.polarity
    tone_analysis = "Positive" if sentiment > 0 else "Negative" if sentiment < 0 else "Neutral"
    print("tone analysed successfully")

    # Analyze English fluency
    try:
        # Count the total number of words in the transcription
        words = len(transcription.split())
        print("hello")
        # Calculate the duration of the audio in minutes
        duration = librosa.get_duration(filename=audio_path) / 60  # Duration in minutes
        print("hello")

        # Calculate words per minute
        if duration > 0:
            fluency_score = words / duration  # Words per minute
        else:
            return "Duration of the audio is zero, unable to calculate fluency."

        # Categorize fluency score
        if fluency_score >= 120:  # 120+ words per minute
            fluency = "Excellent"
        elif 80 <= fluency_score < 120:  # 80 to 119 words per minute
            fluency = "Average"
        else:  # Below 80 words per minute
            fluency = "Poor"

        #return f"Fluency: {fluency} ({fluency_score:.2f} words per minute)"
    except Exception as e:
        return f"Error in analyzing fluency: {e}"

    # Body language analysis
    def analyze_body_language(video_file):
        cap = cv2.VideoCapture(video_file)
        frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        motion_score = 0
        _, prev_frame = cap.read()
        prev_gray = cv2.cvtColor(prev_frame, cv2.COLOR_BGR2GRAY)
        
        for _ in range(frame_count - 1):
            ret, curr_frame = cap.read()
            if not ret:
                break
            curr_gray = cv2.cvtColor(curr_frame, cv2.COLOR_BGR2GRAY)
            diff = cv2.absdiff(prev_gray, curr_gray)
            motion_score += np.sum(diff > 50)
            prev_gray = curr_gray

        cap.release()
        motion_score = motion_score / frame_count
        confidence = "High" if motion_score > 50000 else "Low"
        return confidence

    confidence_analysis = analyze_body_language(video_path)
    print("confidence analysis done successfully")

    # Combine results
    results = {
        "Speech-to-Text": transcription,
        "Tone Analysis": tone_analysis,
        "English Fluency (WPM)": fluency,
        "Confidence Analysis (Body Language)": confidence_analysis
    }
    
    # Save results to a text file
    with open(output_text_path, "w") as f:
        for key, value in results.items():
            f.write(f"{key}: {value}\n")
    
    # Clean up temporary files
    os.remove(audio_path)
    print(f"Analysis for {video_path} saved to {output_text_path}")
    output_folder2='generated_resumes'
    all_paragraphs = process_text_files('analysis_results')
    extract_resume_info(all_paragraphs)
    
    

@app.route('/upload_videos', methods=['POST'])
def upload_videos():
    """
    Endpoint to handle video resume uploads and perform analysis.
    """
    if 'videos' not in request.files:
        return jsonify({"error": "No video files uploaded"}), 400

    uploaded_files = request.files.getlist('videos')
    response_data = []

    for video_file in uploaded_files:
        if video_file.filename == '':
            continue  # Skip empty uploads

        # Save video file
        save_path = os.path.join(app.config['UPLOAD_FOLDER1'], video_file.filename)
        video_file.save(save_path)

        # Perform analysis
        output_text_path = os.path.join(ANALYSIS_FOLDER1, f"{os.path.splitext(video_file.filename)[0]}_analysis.txt")
        process_video(save_path, output_text_path)

        # Append results to response
        response_data.append({
            "video": video_file.filename,
            "analysis": f"Analysis saved to {output_text_path}"
        })

    return jsonify({"message": "Videos processed successfully", "results": response_data}), 200

@app.route('/send_emails', methods=['POST'])
def send_emails():
    # Check if a file is uploaded
    if 'excel_file' not in request.files:
        return "No file part in the request.", 400

    file = request.files['excel_file']

    if file.filename == '':
        return "No file selected.", 400

    try:
        # Read the uploaded Excel file
        df = pd.read_excel(file)
        
        # Extract necessary columns
        if not all(col in df.columns for col in ['Name', 'Email', 'Interview date and time']):
            return "Invalid file format. Make sure the file contains 'Name', 'Email', and 'Interview date and time' columns.", 400
        
        names = df['Name']
        emails = df['Email']
        interview_schedules = df['Interview date and time']

        # Email setup
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(EMAIL_ADDRESS,'umyi lvse tcxp sowv')

        for name, email, schedule in zip(names, emails, interview_schedules):
            subject = "Interview Schedule"
            message = f"Congratulations {name},\n\nYour interview is scheduled on {schedule}.\n\nBest regards,\nHR Team"

            # Create email
            msg = MIMEMultipart()
            msg['From'] = EMAIL_ADDRESS
            msg['To'] = email
            msg['Subject'] = subject
            msg.attach(MIMEText(message, 'plain'))

            # Send email
            server.sendmail(EMAIL_ADDRESS, email, msg.as_string())

        server.quit()

        return "Emails sent successfully."

    except Exception as e:
        print(f"Error: {e}")
        return f"An error occurred: {str(e)}", 500



if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000,debug=True)
